# src/eevee/analysis/data_processor.py

from ..io.readers.input_reader import InputReader
from ..core.thermodynamics import ThermoCorrector
import pandas as pd
from pathlib import Path
from typing import List


def create_gibbs_table(
    input_path: str, 
    temperature: float = 300, 
    fugacity_dict: dict = None
):
    """
    Create a table of Gibbs free energies and thermodynamic corrections
    
    Parameters
    ----------
    input_path : str
        Path to input file
    temperature : float, optional
        Temperature in K, default 300K
    fugacity_dict : dict, optional
        Dictionary of fugacities for gas species (e.g., {'H2': 1e5, 'CO2': 1e5})
        If None, default fugacity of 1e5 Pa is used for all gas species
    
    Returns
    -------
    pd.DataFrame
        Extended DataFrame containing original data and Gibbs free energies
    """
    # Default fugacity dictionary
    if fugacity_dict is None:
        fugacity_dict = {}
    
    # 1) Load input file
    input_reader = InputReader(input_path)
    energy_df = input_reader.energy_df.copy()

    # 2) Calculate free energy for each species using ThermoCorrector
    def calculate_gibbs(row):
        species_name = row['species_name']
        # Get species-specific fugacity
        fugacity = fugacity_dict.get(species_name, 1e5)
        
        thermo_corrector = ThermoCorrector(row)
        status, ZPE, Cp, H, dS, TS, F = thermo_corrector.get_free_energy(temperature, fugacity)
        E = row['formation_energy']
        G = E + F

        # round 3
        return pd.Series({
            'status': status,
            'E': round(E, 3),
            'ZPE': round(ZPE, 3),
            'Cp': round(Cp, 3),
            'H': round(H, 3),
            'dS': round(dS, 3),
            'TS': round(TS, 3),
            'F': round(F, 3),
            'G': round(G, 3)
        })
    
    # Calculate Gibbs corrections and add to energy_df
    gibbs_corrections = energy_df.apply(calculate_gibbs, axis=1)
    extended_df = pd.concat([energy_df, gibbs_corrections], axis=1)
    
    return extended_df

def compare_multiple_gibbs_tables(
    input_paths: List[str],
    temperature: float = 298.15,
    fugacity_dict: dict = None,
    threshold: float = 0.01,
    columns_to_compare: List[str] = ['E', 'F', 'G']
) -> pd.DataFrame:
    """
    Compare Gibbs free energy tables from multiple input files
    
    Parameters
    ----------
    input_paths : List[str]
        List of paths to input files
    temperature : float, optional
        Temperature in K, default 298.15K
    fugacity_dict : dict, optional
        Dictionary of fugacities for gas species, default None
    threshold : float, optional
        Minimum difference to highlight, by default 0.01 eV
    columns_to_compare : List[str], optional
        Columns to compare, by default ['E', 'F', 'G']
        
    Returns
    -------
    pd.DataFrame
        DataFrame containing comparison of energies between input files
    """
    # Store results for each input file
    all_results = {}
    all_species = set()
    
    # Process each input file
    for input_path in input_paths:
        try:
            input_name = Path(input_path).stem
            gibbs_df = create_gibbs_table(input_path, temperature, fugacity_dict)
            
            # Round numeric columns
            numeric_cols = gibbs_df.select_dtypes(include=['float64']).columns
            gibbs_df[numeric_cols] = gibbs_df[numeric_cols].round(2)
            
            # Clean up status column - ensure it's just 'gas' or 'ads' or another simple string
            gibbs_df['status'] = gibbs_df['status'].apply(lambda x: 
                str(x).split()[0] if isinstance(x, str) and ' ' in str(x) else x)
            
            # Convert 'ts' status to 'ads'
            gibbs_df['status'] = gibbs_df['status'].apply(lambda x:
                'ads' if str(x).lower() == 'ts' else x)
            
            all_results[input_name] = gibbs_df
            all_species.update(gibbs_df['species_name'])
            
        except Exception as e:
            raise ValueError(f"Error processing input file {input_path}: {str(e)}")
    
    # Create comparison DataFrame with MultiIndex columns
    comparison_data = []
    
    # Create column MultiIndex tuples
    columns = [
        ('Info', 'species_name'),
        ('Info', 'status')
    ]
    
    # Add columns for each input file (without diff columns)
    for input_path in input_paths:
        input_name = Path(input_path).stem
        for col in columns_to_compare:
            columns.append((input_name, col))
    
    # Add diff columns at the end
    for col in columns_to_compare:
        columns.append(('Diff', col))
    
    # Create MultiIndex columns
    column_index = pd.MultiIndex.from_tuples(
        columns,
        names=['Source', 'Value']
    )
    
    # For each species
    for species in sorted(all_species):
        species_data = {
            ('Info', 'species_name'): species
        }
        
        # Get status - check each file until we find this species
        status = None
        for input_name, df in all_results.items():
            species_row = df[df['species_name'] == species]
            if not species_row.empty:
                # Extract just the simple status value and use the first one
                status_val = species_row.iloc[0]['status']
                if isinstance(status_val, pd.Series):
                    # Get just the first value if it's a Series
                    status = status_val.iloc[0] if not status_val.empty else "unknown"
                    # If it contains additional text, extract just the status type
                    if isinstance(status, str) and ' ' in status:
                        status = status.split()[0]
                else:
                    status = status_val
                    if isinstance(status, str) and ' ' in status:
                        status = status.split()[0]
                
                # Convert 'ts' to 'ads'
                if isinstance(status, str) and status.lower() == 'ts':
                    status = 'ads'
                
                break
                
        species_data[('Info', 'status')] = status if status is not None else "unknown"
                
        # Get values from each input file
        for input_path in input_paths:
            input_name = Path(input_path).stem
            df = all_results[input_name]
            species_row = df[df['species_name'] == species]
            
            # Add values
            for col in columns_to_compare:
                if not species_row.empty:
                    value = species_row[col].iloc[0]
                    # Handle case where value is a Series
                    if isinstance(value, pd.Series):
                        value = value.iloc[0] if not value.empty else None
                    species_data[(input_name, col)] = value
                else:
                    species_data[(input_name, col)] = None
            
        # Calculate differences (second file - first file)
        first_name = Path(input_paths[0]).stem
        second_name = Path(input_paths[1]).stem
        
        for col in columns_to_compare:
            first_val = species_data.get((first_name, col))
            second_val = species_data.get((second_name, col))
            
            if first_val is not None and second_val is not None:
                diff = second_val - first_val
                species_data[('Diff', col)] = diff
            else:
                species_data[('Diff', col)] = None
        
        comparison_data.append(species_data)
    
    # Create DataFrame with MultiIndex columns
    comparison_df = pd.DataFrame(comparison_data, columns=column_index)
    
    # Sort by status (gas first) and then by species name
    # First create a temporary sorting column
    comparison_df['_sort_status'] = comparison_df[('Info', 'status')].apply(
        lambda x: 0 if str(x).lower() == 'gas' else 1
    )
    
    # Sort the DataFrame
    comparison_df = comparison_df.sort_values(
        by=['_sort_status', ('Info', 'species_name')]
    )
    
    # Remove the temporary column without triggering the warning
    comparison_df = comparison_df[column_index]
    
    # Round numeric columns to 2 decimal places
    numeric_cols = comparison_df.select_dtypes(include=['float64']).columns
    comparison_df[numeric_cols] = comparison_df[numeric_cols].round(2)
    
    # Add metadata for styling
    comparison_df.attrs['threshold'] = threshold
    
    return comparison_df

def style_differences(df: pd.DataFrame, hide_index_names: bool = True) -> pd.DataFrame:
    """
    Create styled DataFrame with highlighted differences
    
    Parameters
    ----------
    df : pd.DataFrame
        DataFrame from compare_multiple_gibbs_tables
    hide_index_names : bool, optional
        Whether to hide the MultiIndex names (Source, Value), default True
        
    Returns
    -------
    pd.DataFrame
        Styled DataFrame with highlighted differences
    """
    # Make a copy to avoid modifying the original
    styled_df = df.copy()
    
    # Get threshold from DataFrame attributes
    threshold = styled_df.attrs.get('threshold', 0.01)
    
    # Define function to color cells
    def highlight_diff(val, col):
        # Check if this is a difference column (under 'Diff' source)
        if not isinstance(col, tuple) or col[0] != 'Diff':
            return ''
        
        # Skip non-numeric values
        if not isinstance(val, (int, float)) or pd.isna(val):
            return ''
        
        # Apply coloring based on magnitude and sign
        if abs(val) < threshold:
            return ''
        elif val > 0:
            # Positive difference (red)
            return f'background-color: rgba(255, 0, 0, {min(abs(val)/0.5, 0.5)})'
        else:
            # Negative difference (green)
            return f'background-color: rgba(0, 128, 0, {min(abs(val)/0.5, 0.5)})'
    
    # Pre-process any remaining Series objects in the DataFrame
    for col in styled_df.columns:
        styled_df[col] = styled_df[col].apply(
            lambda x: x.iloc[0] if isinstance(x, pd.Series) and not x.empty else x
        )
    
    # Apply style and format numbers to 2 decimal places
    styled = styled_df.style.apply(
        lambda row: [highlight_diff(val, col) for col, val in row.items()], 
        axis=1
    )
    
    # Format numeric columns to 2 decimal places
    numeric_cols = styled_df.select_dtypes(include=['float64']).columns
    styled = styled.format("{:.2f}", subset=numeric_cols)
    
    # Hide MultiIndex names if requested
    if hide_index_names:
        styled = styled.hide(names=False)
    
    return styled

def style_mechanism_differences(df: pd.DataFrame) -> pd.DataFrame:
    """
    스타일이 적용된 DataFrame을 반환합니다.
    
    Parameters
    ----------
    df : pd.DataFrame
        compare_mechanism_differences()에서 반환된 DataFrame
        
    Returns
    -------
    pd.DataFrame
        스타일이 적용된 DataFrame
    """
    threshold = df.attrs.get('threshold', 0.05)
    
    def highlight_diff(val, col):
        """단일 값에 대한 스타일링"""
        # Check if this is a difference column (under 'Diff' source)
        if col[0] == 'Diff':
            if pd.isna(val):
                return None
            if abs(val) > threshold:
                return 'background-color: #ffcdd2' if val > 0 else 'background-color: #c8e6c9'
        return None
    
    # Format numbers and apply highlighting
    styled = df.style.apply(lambda x: [highlight_diff(v, c) for v, c in zip(x, df.columns)], axis=1)
    
    # Format numeric columns to 2 decimal places
    numeric_cols = df.select_dtypes(include=['float64']).columns
    styled = styled.format("{:.2f}", subset=numeric_cols)
    
    return styled

def compare_multiple_inputs(
    input_paths: List[str],
    mkm_path: str,
    temperature: float = 300.0,
    potential: float = 0.0,
    threshold: float = 0.05
) -> pd.DataFrame:
    """
    Compare reaction mechanisms for multiple input files using the same MKM file
    
    Parameters
    ----------
    input_paths : list[str]
        List of paths to input files
    mkm_path : str
        Path to MKM file
    temperature : float, optional
        Temperature in K, default 300K
    potential : float, optional
        Applied potential (V vs. RHE), default 0.0V
    threshold : float, optional
        Minimum difference to highlight, by default 0.05 eV
        
    Returns
    -------
    pd.DataFrame
        DataFrame containing comparison of dG and Ga for each input
    """
    from .mechanism import MechanismAnalyzer
    
    # Store results for each input file
    all_results = {}
    
    # Initialize analyzers with specified temperature
    for input_path in input_paths:
        input_name = Path(input_path).stem
        analyzer = MechanismAnalyzer(
            mkm_path=mkm_path,
            input_path=input_path,
            temperature=temperature  # Pass temperature to analyzer
        )
        
        # Get mechanism data
        df = analyzer.mechanism_to_dataframe(potential=potential)
        all_results[input_name] = df
    
    # Create comparison DataFrame
    comparison_data = []
    
    # Get reference input (first one)
    ref_input = list(all_results.keys())[0]
    ref_df = all_results[ref_input]
    
    # For each reaction step
    for idx, ref_row in ref_df.iterrows():
        step_data = {
            'step_idx': ref_row['step_idx'],
            'rxn_idx': ref_row['rxn_idx'],
            'Reaction': ref_row['Reaction'],
            f'{ref_input}_dG': ref_row['dG'],
            f'{ref_input}_Ga': ref_row['Ga'],
        }
        
        # Add data from other inputs
        for input_name, df in all_results.items():
            if input_name == ref_input:
                continue
                
            row = df[df['rxn_idx'] == ref_row['rxn_idx']].iloc[0]
            step_data[f'{input_name}_dG'] = row['dG']
            step_data[f'{input_name}_Ga'] = row['Ga']
            
            # Calculate differences
            step_data[f'{input_name}_dG_diff'] = row['dG'] - ref_row['dG']
            step_data[f'{input_name}_Ga_diff'] = row['Ga'] - ref_row['Ga']
        
        comparison_data.append(step_data)
    
    comparison_df = pd.DataFrame(comparison_data)
    
    # Round numeric columns to 2 decimal places
    numeric_cols = comparison_df.select_dtypes(include=['float64']).columns
    comparison_df[numeric_cols] = comparison_df[numeric_cols].round(2)
    
    # Add metadata for styling
    comparison_df.attrs['threshold'] = threshold
    
    return comparison_df

def save_mechanism_comparison(
    df: pd.DataFrame,
    output_path: str,
    threshold: float = 0.05,
    index: bool = False
) -> None:
    """
    스타일이 적용된 메커니즘 비교 테이블을 HTML 파일로 저장합니다.
    
    Parameters
    ----------
    df : pd.DataFrame
        비교할 DataFrame
    output_path : str
        저장할 HTML 파일 경로
    threshold : float, optional
        강조 표시할 차이 값의 임계값, 기본값 0.05 eV
    index : bool, optional
        인덱스 포함 여부, 기본값 False
    """
    styled_df = style_mechanism_differences(df)
    
    # Convert output_path to Path object
    output_path = Path(output_path)
    
    # Ensure the path has .html extension
    if output_path.suffix != '.html':
        output_path = output_path.with_suffix('.html')
        
    # Create directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Save to HTML
    if not index:
        styled_df = styled_df.hide(axis='index')
    
    styled_df.to_html(output_path)

def compare_temperatures(
    input_path: str,
    temperatures: List[float],
    fugacity: float = 1e5,
    threshold: float = 0.01,
    columns_to_compare: List[str] = ['E', 'F', 'G']
) -> pd.DataFrame:
    """
    Compare Gibbs free energy tables at different temperatures
    
    Parameters
    ----------
    input_path : str
        Path to input file
    temperatures : List[float]
        List of temperatures in K to compare
    fugacity : float, optional
        Fugacity in Pa, by default 1e5
    threshold : float, optional
        Minimum difference to highlight, by default 0.01 eV
    columns_to_compare : List[str], optional
        Columns to compare, by default ['E', 'F', 'G']
        
    Returns
    -------
    pd.DataFrame
        DataFrame comparing energies at different temperatures
    """
    # Calculate Gibbs energies for each temperature
    all_results = {}
    for temp in temperatures:
        gibbs_df = create_gibbs_table(
            input_path=input_path,
            temperature=temp,
            fugacity=fugacity
        )
        all_results[f"{temp}K"] = gibbs_df
        
    # Get all unique species
    all_species = set()
    for df in all_results.values():
        all_species.update(df['species_name'])
    
    # Prepare comparison data
    comparison_data = []
    ref_temp = temperatures[0]
    ref_name = f"{ref_temp}K"
    
    # For each species
    for species in sorted(all_species):
        species_data = {'species_name': species}
        
        # Get status from first temperature result
        for df in all_results.values():
            species_row = df[df['species_name'] == species]
            if not species_row.empty:
                species_data['status'] = species_row.iloc[0]['status']
                break
        
        # For each column to compare
        for col in columns_to_compare:
            # Add values from each temperature
            for temp_name, df in all_results.items():
                species_row = df[df['species_name'] == species]
                value = species_row[col].iloc[0] if not species_row.empty else None
                species_data[f"{temp_name}_{col}"] = value
                
                # Calculate differences from reference temperature
                if temp_name != ref_name:
                    ref_row = all_results[ref_name][all_results[ref_name]['species_name'] == species]
                    if not ref_row.empty and not species_row.empty:
                        diff = value - ref_row[col].iloc[0]
                        species_data[f"{temp_name}_{col}_diff"] = diff
        
        comparison_data.append(species_data)
    
    comparison_df = pd.DataFrame(comparison_data)
    
    # Add metadata for styling
    comparison_df.attrs['columns_to_compare'] = columns_to_compare
    comparison_df.attrs['threshold'] = threshold
    comparison_df.attrs['reference_temp'] = ref_temp
    
    # Round numeric columns
    numeric_cols = comparison_df.select_dtypes(include=['float64']).columns
    comparison_df[numeric_cols] = comparison_df[numeric_cols].round(2)
    
    return comparison_df

def style_temperature_differences(df: pd.DataFrame) -> pd.DataFrame:
    """
    Create styled DataFrame with highlighted temperature differences
    
    Parameters
    ----------
    df : pd.DataFrame
        DataFrame from compare_temperatures
        
    Returns
    -------
    pd.DataFrame
        Styled DataFrame with highlighted differences
    """
    if 'columns_to_compare' not in df.attrs:
        raise ValueError("DataFrame must be from compare_temperatures()")
        
    threshold = df.attrs.get('threshold', 0.01)
    
    def highlight_diff(val, col):
        if not col.endswith('_diff'):
            return None
        if pd.isna(val):
            return None
        if abs(val) > threshold:
            if val > 0:
                return 'background-color: #ffcdd2'  # Red (positive)
            else:
                return 'background-color: #c8e6c9'  # Green (negative)
        return None
    
    # Format numbers and apply highlighting
    styled = df.style.apply(lambda x: [highlight_diff(v, c) for v, c in zip(x, df.columns)], axis=1)
    
    # Format numeric columns to 3 decimal places
    numeric_cols = df.select_dtypes(include=['float64']).columns
    styled = styled.format("{:.3f}", subset=numeric_cols)
    
    return styled

def save_temperature_comparison(
    df: pd.DataFrame,
    output_path: str,
    threshold: float = 0.01,
    index: bool = False
) -> None:
    """
    Save styled temperature comparison table to HTML file
    
    Parameters
    ----------
    df : pd.DataFrame
        DataFrame from compare_temperatures
    output_path : str
        Path to save the HTML file
    threshold : float, optional
        Minimum difference to highlight, by default 0.01 eV
    index : bool, optional
        Whether to include index in output, by default False
    """
    styled_df = style_temperature_differences(df)
    
    # Convert output_path to Path object
    output_path = Path(output_path)
    
    # Ensure the path has .html extension
    if output_path.suffix != '.html':
        output_path = output_path.with_suffix('.html')
        
    # Create directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Save to HTML
    styled_df.hide(axis='index').to_html(output_path)

def compare_mechanism_temperatures(
    input_path: str,
    mkm_path: str,
    temperatures: List[float],
    potential: float = 0.0,
    threshold: float = 0.05
) -> pd.DataFrame:
    """
    Compare reaction mechanisms at different temperatures
    
    Parameters
    ----------
    input_path : str
        Path to input file
    mkm_path : str
        Path to MKM file
    temperatures : List[float]
        List of temperatures in K to compare
    potential : float, optional
        Applied potential (V vs. RHE), default 0.0V
    threshold : float, optional
        Minimum difference to highlight, by default 0.05 eV
        
    Returns
    -------
    pd.DataFrame
        DataFrame containing mechanism comparison at different temperatures
    """
    from .mechanism import MechanismAnalyzer
    
    # Store results for each temperature
    all_results = {}
    
    # Analyze mechanism at each temperature
    for temp in temperatures:
        print(f"\nAnalyzing temperature: {temp}K")
        analyzer = MechanismAnalyzer(
            mkm_path=mkm_path,
            input_path=input_path,
            temperature=temp
        )
        df = analyzer.mechanism_to_dataframe(potential=potential)
        all_results[f"{temp}K"] = df
    
    # Create comparison DataFrame with MultiIndex columns
    comparison_data = []
    ref_temp = temperatures[0]
    ref_name = f"{ref_temp}K"
    ref_df = all_results[ref_name]
    
    # Create column MultiIndex tuples
    columns = [
        ('Info', 'step_idx'),
        ('Info', 'rxn_idx'),
        ('Info', 'Reaction')
    ]
    
    for temp in temperatures:
        temp_name = f"{temp}K"
        # Add base columns
        columns.extend([
            (temp_name, 'dG'),
            (temp_name, 'Ga'),
            (temp_name, 'IS')
        ])
        # Add diff columns for non-reference temperatures
        if temp != ref_temp:
            columns.extend([
                (temp_name, 'dG_diff'),
                (temp_name, 'Ga_diff'),
                (temp_name, 'IS_diff')
            ])
    
    # Create MultiIndex columns
    column_index = pd.MultiIndex.from_tuples(
        columns,
        names=['Temperature', 'Value']
    )
    
    # For each reaction step
    for idx, ref_row in ref_df.iterrows():
        row_data = {
            ('Info', 'step_idx'): ref_row['step_idx'],
            ('Info', 'rxn_idx'): ref_row['rxn_idx'],
            ('Info', 'Reaction'): ref_row['Reaction']
        }
        
        # Add data for each temperature
        for temp in temperatures:
            temp_name = f"{temp}K"
            df = all_results[temp_name]
            row = df[df['rxn_idx'] == ref_row['rxn_idx']].iloc[0]
            
            # Base values
            row_data[(temp_name, 'dG')] = row['dG']
            row_data[(temp_name, 'Ga')] = row['Ga']
            row_data[(temp_name, 'IS')] = row['IS']
            
            # Difference values (for non-reference temperatures)
            if temp != ref_temp:
                row_data[(temp_name, 'dG_diff')] = row['dG'] - ref_row['dG']
                row_data[(temp_name, 'Ga_diff')] = row['Ga'] - ref_row['Ga']
                row_data[(temp_name, 'IS_diff')] = row['IS'] - ref_row['IS']
        
        comparison_data.append(row_data)
    
    # Create DataFrame with MultiIndex columns
    comparison_df = pd.DataFrame(comparison_data, columns=column_index)
    
    # Round numeric columns
    numeric_cols = comparison_df.select_dtypes(include=['float64']).columns
    comparison_df[numeric_cols] = comparison_df[numeric_cols].round(2)
    
    # Add metadata for styling
    comparison_df.attrs['threshold'] = threshold
    comparison_df.attrs['reference_temp'] = ref_temp
    
    return comparison_df

def style_mechanism_temperature_differences(df: pd.DataFrame) -> pd.DataFrame:
    threshold = df.attrs.get('threshold', 0.05)
    
    def highlight_diff(val):
        """단일 값에 대한 스타일링"""
        try:
            # MultiIndex에서 현재 열 이름 가져오기
            col_name = highlight_diff.col_name
            if not col_name[1].endswith('_diff'):
                return None
            if pd.isna(val):
                return None
            if abs(val) > threshold:
                if val > 0:
                    return 'background-color: #ffcdd2'  # Red (positive)
                else:
                    return 'background-color: #c8e6c9'  # Green (negative)
        except:
            return None
        return None
    
    # Format numbers and apply highlighting
    styled = df.style
    
    # 각 열에 대해 개별적으로 스타일 적용
    for col in df.columns:
        if col[1].endswith('_diff'):
            highlight_diff.col_name = col  # 현재 열 이름 저장
            styled = styled.applymap(highlight_diff, subset=[col])
    
    # Format numeric columns to 2 decimal places
    numeric_cols = df.select_dtypes(include=['float64']).columns
    styled = styled.format("{:.2f}", subset=numeric_cols)
    
    return styled

def save_mechanism_temperature_comparison(
    df: pd.DataFrame,
    output_path: str,
    threshold: float = 0.05,
    index: bool = False
) -> None:
    """
    Save styled mechanism temperature comparison table to HTML file
    
    Parameters
    ----------
    df : pd.DataFrame
        DataFrame from compare_mechanism_temperatures
    output_path : str
        Path to save the HTML file
    threshold : float, optional
        Minimum difference to highlight, by default 0.05 eV
    index : bool, optional
        Whether to include index in output, by default False
    """
    styled_df = style_mechanism_temperature_differences(df)
    
    # Convert output_path to Path object
    output_path = Path(output_path)
    
    # Ensure the path has .html extension
    if output_path.suffix != '.html':
        output_path = output_path.with_suffix('.html')
        
    # Create directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Save to HTML
    styled_df.hide(axis='index').to_html(output_path)

def compare_mechanism_differences(
    input_paths: List[str],
    mkm_path: str,
    temperature: float = 300,
    potential: float = 0.0,
    threshold: float = 0.05
) -> pd.DataFrame:
    """메커니즘 비교 함수"""
    from .mechanism import MechanismAnalyzer
    
    if not input_paths:
        raise ValueError("At least one input path must be provided")
    
    # 기준 메커니즘 분석
    base_path = input_paths[0]
    base_name = Path(base_path).stem
    base_analyzer = MechanismAnalyzer(
        mkm_path=mkm_path,
        input_path=base_path,
        temperature=temperature
    )
    base_df = base_analyzer.mechanism_to_dataframe(potential=potential)
    
    # Create column MultiIndex tuples
    columns = [
        ('Info', 'step_idx'),
        ('Info', 'rxn_idx'),
        ('Info', 'Reaction')
    ]
    
    # Add columns for base file
    columns.extend([
        (base_name, 'dG'),
        (base_name, 'Ga')
    ])
    
    # 비교할 메커니즘들 분석
    all_results = {}
    for comp_path in input_paths[1:]:  # Skip the first (base) path
        name = Path(comp_path).stem
        analyzer = MechanismAnalyzer(
            mkm_path=mkm_path,
            input_path=comp_path,
            temperature=temperature
        )
        df = analyzer.mechanism_to_dataframe(potential=potential)
        all_results[name] = df
        
        # Add columns for comparison files
        columns.extend([
            (name, 'dG'),
            (name, 'Ga')
        ])
    
    # Add diff columns at the end
    columns.extend([
        ('Diff', 'dG'),
        ('Diff', 'Ga')
    ])
    
    # Create MultiIndex columns
    column_index = pd.MultiIndex.from_tuples(
        columns,
        names=['Source', 'Value']
    )
    
    # Prepare comparison data
    comparison_data = []
    for idx, base_row in base_df.iterrows():
        row_data = {
            ('Info', 'step_idx'): base_row['step_idx'],
            ('Info', 'rxn_idx'): base_row['rxn_idx'],
            ('Info', 'Reaction'): base_row['Reaction'],
            (base_name, 'dG'): base_row['dG'],
            (base_name, 'Ga'): base_row['Ga']
        }
        
        # Compare with other mechanisms
        for name, comp_df in all_results.items():
            comp_row = comp_df[comp_df['rxn_idx'] == base_row['rxn_idx']].iloc[0]
            
            row_data[(name, 'dG')] = comp_row['dG']
            row_data[(name, 'Ga')] = comp_row['Ga']
            
            # Store differences under 'Diff' source
            if name == Path(input_paths[1]).stem:  # Only for the first comparison file
                row_data[('Diff', 'dG')] = comp_row['dG'] - base_row['dG']
                row_data[('Diff', 'Ga')] = comp_row['Ga'] - base_row['Ga']
        
        comparison_data.append(row_data)
    
    # Create DataFrame with MultiIndex columns
    comparison_df = pd.DataFrame(comparison_data, columns=column_index)
    
    # Round numeric columns
    numeric_cols = comparison_df.select_dtypes(include=['float64']).columns
    comparison_df[numeric_cols] = comparison_df[numeric_cols].round(2)
    
    # Add metadata for styling
    comparison_df.attrs['threshold'] = threshold
    comparison_df.attrs['base_name'] = base_name
    
    return comparison_df

def create_input_from_gibbs(
    extended_df_dl: pd.DataFrame,
    extended_df_gas: pd.DataFrame,
    dl_temperature: float,
    gas_temperature: float,
    output_path: str = None
) -> pd.DataFrame:
    """
    Create input format DataFrame using Gibbs free energies from two temperature corrections.

    For adsorbed species (i.e. rows whose site_name does not include 'gas'),
    Gibbs energies are taken from extended_df_dl computed with the dl_temperature
    (cathode temperature). For gas-phase species (i.e. rows whose site_name includes 'gas'),
    Gibbs energies are calculated using gas_temperature (electrolyte temperature).
    In addition, gas-phase rows are duplicated so that one copy retains the original site_name
    ('gas') while the duplicate has its site_name replaced with 'dl'.

    Parameters
    ----------
    extended_df_dl : pd.DataFrame
        Extended DataFrame from create_gibbs_table() computed with dl_temperature.
    extended_df_gas : pd.DataFrame
        Extended DataFrame from create_gibbs_table() computed with gas_temperature.
    dl_temperature : float
        Temperature used for dl calculations (cathode temperature).
    gas_temperature : float
        Temperature used for gas calculations (electrolyte temperature).
    output_path : str, optional
        Path to save the output file. If None, file is not saved.
        
    Returns
    -------
    pd.DataFrame
        Input format DataFrame with duplicated gas-phase rows.
    """
    # Process adsorbed species from extended_df_dl (i.e. not gas-phase)
    df_ads = extended_df_dl[~extended_df_dl['site_name'].str.contains('gas', case=False, na=False)].copy()
    df_gas_dl = extended_df_dl[extended_df_dl['site_name'].str.contains('gas', case=False, na=False)].copy()
    df_gas_gas = extended_df_gas[extended_df_gas['site_name'].str.contains('gas', case=False, na=False)].copy()


    # 가스 phase 항목들은 surface_name을 'None'으로 설정 (출력 예와 맞추기 위함).
    df_gas_gas.loc[:, 'surface_name'] = 'None'
    df_gas_dl.loc[:, 'surface_name'] = 'Cu'


    # Duplicate the gas-phase rows:
    #  - one copy retains the original site_name ("gas")
    #  - the duplicate changes site_name to "dl"
    
    df_gas_dl.loc[:, 'site_name'] = 'dl'

    # Combine adsorbed species and both sets of gas-phase rows
    combined_df = pd.concat([df_gas_gas, df_gas_dl, df_ads], ignore_index=True)

    # Build the output DataFrame by selecting the necessary columns
    output_df = combined_df[['surface_name', 'site_name', 'species_name', 'frequencies', 'reference']].copy()
    # Use the free energy column (G) as the formation energy
    output_df['formation_energy'] = combined_df['G']

    # If an output file is requested, convert frequencies to a string and save the file
    if output_path:
        output_df['frequencies'] = output_df['frequencies'].apply(lambda x: '[]')
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_df.to_csv(output_path, sep='\t', index=False, na_rep="None")

    return output_df

def compare_gibbs_tables(
    input_paths: List[str],
    temperature: float = 298.15,
    fugacity_dict: dict = None,
    threshold: float = 0.01,
    columns_to_compare: List[str] = ['formation_energy', 'F', 'G']
) -> pd.DataFrame:
    """
    Compare Gibbs tables from multiple input files.
    
    Args:
        input_paths: List of paths to the input files to compare
        temperature: Temperature in K
        fugacity_dict: Dictionary of species-specific fugacities (Pa)
        threshold: Threshold for highlighting differences
        columns_to_compare: Columns to compare
        
    Returns:
        DataFrame with comparison of Gibbs tables
    """
    if fugacity_dict is None:
        # Default fugacity is 1 bar = 100,000 Pa
        fugacity_dict = {}
        
    all_dfs = []
    all_species = set()
    
    # Process each input file
    for i, input_path in enumerate(input_paths):
        # Extract name from path
        name = Path(input_path).stem
        
        try:
            # Calculate Gibbs energies
            gibbs_df = create_gibbs_table(input_path, temperature, fugacity_dict)
            
            # Clean up status column - ensure it's just 'gas' or 'ads' or another simple string
            gibbs_df['status'] = gibbs_df['status'].apply(lambda x: 
                str(x).split()[0] if isinstance(x, str) and ' ' in str(x) else x)
            
            # Convert 'ts' status to 'ads'
            gibbs_df['status'] = gibbs_df['status'].apply(lambda x:
                'ads' if str(x).lower() == 'ts' else x)
            
            # Add source column for easy identification
            gibbs_df['Source'] = name
            
            all_dfs.append(gibbs_df)
            
            # Update set of all species
            species_ids = set(gibbs_df['species_id'])
            all_species.update(species_ids)
            
        except Exception as e:
            raise ValueError(f"Error processing input file {input_path}: {str(e)}")
    
    # Create comparison DataFrame with MultiIndex columns
    comparison_data = []
    
    # Create column MultiIndex tuples
    columns = [
        ('Info', 'surface_name'),
        ('Info', 'site_name'),
        ('Info', 'species_name'),
        ('Info', 'status')
    ]
    
    # Add columns for each input file
    for input_path in input_paths:
        input_name = Path(input_path).stem
        for col in columns_to_compare:
            columns.append((input_name, col))
    
    # Add diff columns at the end (comparing second file to first file)
    for col in columns_to_compare:
        columns.append(('Diff', col))
    
    # Create MultiIndex columns
    column_index = pd.MultiIndex.from_tuples(
        columns,
        names=['Source', 'Value']
    )
    
    # For each unique species combination
    for species_id in sorted(all_species):
        surface_name, site_name, species_name = species_id.split('_')
        species_data = {
            ('Info', 'surface_name'): surface_name,
            ('Info', 'site_name'): site_name,
            ('Info', 'species_name'): species_name
        }
        
        # Get status from first file that has this species
        status = None
        for df in all_dfs:
            mask = (
                (df['surface_name'] == surface_name) & 
                (df['site_name'] == site_name) & 
                (df['species_name'] == species_name)
            )
            if mask.any():
                status_val = df[mask].iloc[0]['status']
                if isinstance(status_val, pd.Series):
                    # Handle Series objects
                    status = status_val.iloc[0] if not status_val.empty else "unknown"
                    # Clean up status string if needed
                    if isinstance(status, str) and ' ' in status:
                        status = status.split()[0]
                else:
                    status = status_val
                    if isinstance(status, str) and ' ' in status:
                        status = status.split()[0]
                
                # Convert 'ts' to 'ads'
                if isinstance(status, str) and status.lower() == 'ts':
                    status = 'ads'
                    
                break
        
        species_data[('Info', 'status')] = status if status is not None else "unknown"
        
        # Get values from each input file
        first_values = {}  # Store first file's values for diff calculation
        for i, input_path in enumerate(input_paths):
            input_name = Path(input_path).stem
            df = all_dfs[i]
            
            mask = (
                (df['surface_name'] == surface_name) & 
                (df['site_name'] == site_name) & 
                (df['species_name'] == species_name)
            )
            row = df[mask]
            
            # Add values
            for col in columns_to_compare:
                if not row.empty:
                    value = row[col].iloc[0]
                    # Handle case where value is a Series
                    if isinstance(value, pd.Series):
                        value = value.iloc[0] if not value.empty else None
                    species_data[(input_name, col)] = value
                    
                    # Store first file's values for diff calculation
                    if i == 0:
                        first_values[col] = value
                else:
                    species_data[(input_name, col)] = None
        
        # Calculate differences (second file - first file)
        if len(input_paths) >= 2:
            second_name = Path(input_paths[1]).stem
            for col in columns_to_compare:
                first_val = first_values.get(col)
                second_val = species_data.get((second_name, col))
                
                if first_val is not None and second_val is not None:
                    diff = second_val - first_val
                    species_data[('Diff', col)] = diff
                else:
                    species_data[('Diff', col)] = None
        
        comparison_data.append(species_data)
    
    # Create DataFrame with MultiIndex columns
    comparison_df = pd.DataFrame(comparison_data, columns=column_index)
    
    # Round numeric columns to 2 decimal places
    numeric_cols = comparison_df.select_dtypes(include=['float64']).columns
    comparison_df[numeric_cols] = comparison_df[numeric_cols].round(2)
    
    # Add metadata for styling
    comparison_df.attrs['threshold'] = threshold
    
    return comparison_df


def save_gibbs_comparison(
    df: pd.DataFrame,
    output_path: str,
    threshold: float = 0.01,
    index: bool = False
) -> None:
    """
    Save styled Gibbs comparison table to HTML file
    
    Parameters
    ----------
    df : pd.DataFrame
        DataFrame from compare_multiple_gibbs_tables
    output_path : str
        Path to save the HTML file
    threshold : float, optional
        Minimum difference to highlight, by default 0.01 eV
    index : bool, optional
        Whether to include index in output, by default False
    """
    styled_df = style_differences(df)
    
    # Convert output_path to Path object
    output_path = Path(output_path)
    
    # Ensure the path has .html extension
    if output_path.suffix != '.html':
        output_path = output_path.with_suffix('.html')
        
    # Create directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Save to HTML
    if not index:
        styled_df = styled_df.hide(axis='index')
    
    styled_df.to_html(output_path)

def plot_energy_differences(comparison_df, column='G', figsize=(10, 6)):
    """
    Plot histogram of energy differences between files
    
    Parameters
    ----------
    comparison_df : pd.DataFrame
        DataFrame from compare_multiple_gibbs_tables
    column : str, optional
        Column to plot differences for ('E', 'F', or 'G'), by default 'G'
    figsize : tuple, optional
        Figure size, by default (10, 6)
    """
    import matplotlib.pyplot as plt
    import numpy as np
    
    # Extract difference values for the specified column
    diff_values = comparison_df[('Diff', column)].dropna().values
    
    # Create figure
    fig, ax = plt.subplots(figsize=figsize)
    
    # Plot histogram
    bins = np.linspace(min(diff_values) - 0.05, max(diff_values) + 0.05, 30)
    ax.hist(diff_values, bins=bins, alpha=0.7, color='skyblue', edgecolor='black')
    
    # Add vertical line at zero
    ax.axvline(x=0, color='grey', linestyle='--', alpha=0.7)
    
    # Add statistics annotation
    mean_diff = np.mean(diff_values)
    std_diff = np.std(diff_values)
    max_abs_diff = np.max(np.abs(diff_values))
    
    stats_text = (
        f"Mean diff: {mean_diff:.3f} eV\n"
        f"Std dev: {std_diff:.3f} eV\n"
        f"Max abs diff: {max_abs_diff:.3f} eV"
    )
    ax.text(0.95, 0.95, stats_text, transform=ax.transAxes, 
            verticalalignment='top', horizontalalignment='right',
            bbox=dict(boxstyle='round', facecolor='white', alpha=0.7))
    
    # Set labels and title
    ax.set_xlabel(f'Difference in {column} (eV)')
    ax.set_ylabel('Frequency')
    ax.set_title(f'Distribution of {column} energy differences')
    
    # Show grid
    ax.grid(alpha=0.3)
    
    plt.tight_layout()
    return fig, ax

def plot_energy_heatmap(comparison_df, column='G', figsize=(12, 10)):
    """
    Create a heatmap showing energy values across multiple files
    
    Parameters
    ----------
    comparison_df : pd.DataFrame
        DataFrame from compare_multiple_gibbs_tables
    column : str, optional
        Column to compare ('E', 'F', or 'G'), by default 'G'
    figsize : tuple, optional
        Figure size, by default (12, 10)
    """
    import matplotlib.pyplot as plt
    import seaborn as sns
    import pandas as pd
    import numpy as np
    
    # Extract file names (excluding 'Info' and 'Diff')
    file_names = [col[0] for col in comparison_df.columns 
                 if col[0] not in ['Info', 'Diff']]
    
    # Create a pivot table for the heatmap
    # Species as rows, files as columns
    heatmap_data = pd.DataFrame()
    
    for file_name in file_names:
        values = comparison_df[(file_name, column)]
        heatmap_data[file_name] = values
    
    # Set species names as index
    heatmap_data.index = comparison_df[('Info', 'species_name')]
    
    # Drop rows with any NaN values
    heatmap_data = heatmap_data.dropna()
    
    # Create pairwise difference matrix if there are more than 2 files
    if len(file_names) > 2:
        plt.figure(figsize=figsize)
        
        # Calculate pairwise differences
        diff_matrix = pd.DataFrame(index=file_names, columns=file_names)
        
        for i, file1 in enumerate(file_names):
            for j, file2 in enumerate(file_names):
                if i == j:
                    diff_matrix.loc[file1, file2] = 0.0
                else:
                    # Calculate mean absolute difference for non-NaN values
                    valid_rows = ~(heatmap_data[file1].isna() | heatmap_data[file2].isna())
                    if valid_rows.sum() > 0:
                        mean_abs_diff = np.mean(np.abs(
                            heatmap_data.loc[valid_rows, file1] - 
                            heatmap_data.loc[valid_rows, file2]
                        ))
                        diff_matrix.loc[file1, file2] = float(mean_abs_diff)
                    else:
                        diff_matrix.loc[file1, file2] = np.nan
        
        # Ensure all values are numeric by explicitly converting to float
        diff_matrix = diff_matrix.astype(float)
        
        # Plot the difference matrix
        ax = plt.subplot(111)
        sns.heatmap(diff_matrix, annot=True, fmt=".3f", cmap="YlOrRd", ax=ax)
        ax.set_title(f'Mean Absolute Difference in {column} (eV)')
        
        return plt.gcf(), ax
    
    # For 2 or fewer files, show the actual values
    else:
        fig, ax = plt.subplots(figsize=figsize)
        # Ensure heatmap_data is numeric
        heatmap_data = heatmap_data.astype(float)
        sns.heatmap(heatmap_data, annot=True, fmt=".2f", cmap="YlGnBu", ax=ax)
        ax.set_title(f'{column} Values (eV) by Species and Input File')
        
        return fig, ax
    
def create_interactive_explorer(comparison_df, output_path=None):
    """
    Create an interactive Plotly dashboard to explore energy differences
    
    Parameters
    ----------
    comparison_df : pd.DataFrame
        DataFrame from compare_multiple_gibbs_tables
    output_path : str, optional
        Path to save HTML file, by default None
        
    Returns
    -------
    plotly.graph_objects.Figure
        Interactive figure
    """
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    import pandas as pd
    import numpy as np
    
    # Extract file names (excluding 'Info' and 'Diff')
    file_names = [col[0] for col in comparison_df.columns 
                 if col[0] not in ['Info', 'Diff']]
    
    # Create a clean dataframe for plotting
    plot_df = pd.DataFrame({
        'Species': comparison_df[('Info', 'species_name')],
        'Status': comparison_df[('Info', 'status')]
    })
    
    # Add energy values for each file
    for energy_type in ['E', 'F', 'G']:
        for file_name in file_names:
            plot_df[f"{file_name}_{energy_type}"] = comparison_df[(file_name, energy_type)]
    
    # Calculate differences if there are exactly 2 files
    if len(file_names) == 2:
        for energy_type in ['E', 'F', 'G']:
            plot_df[f"Diff_{energy_type}"] = comparison_df[('Diff', energy_type)]
    
    # Drop rows with all NaN energy values
    energy_cols = [col for col in plot_df.columns if col not in ['Species', 'Status']]
    plot_df = plot_df.dropna(subset=energy_cols, how='all')
    
    # Create interactive plot
    fig = make_subplots(
        rows=3, cols=1,
        subplot_titles=('Formation Energy (E)', 'Free Energy Correction (F)', 'Gibbs Energy (G)'),
        shared_xaxes=True,
        vertical_spacing=0.1
    )
    
    # Energy types
    energy_types = ['E', 'F', 'G']
    
    # Color by status
    status_colors = {
        'gas': 'rgba(255, 0, 0, 0.7)',  # Red for gas
        'ads': 'rgba(0, 0, 255, 0.7)'   # Blue for ads
    }
    default_color = 'rgba(100, 100, 100, 0.7)'  # Gray for others
    
    # Add traces for each file and energy type
    for i, energy_type in enumerate(energy_types):
        row = i + 1
        
        # Symbol styles for different files
        symbols = ['circle', 'square', 'diamond', 'cross', 'x', 'star']
        
        for j, file_name in enumerate(file_names):
            # Get colors based on status
            colors = [status_colors.get(str(s).lower(), default_color) for s in plot_df['Status']]
            
            fig.add_trace(
                go.Scatter(
                    x=plot_df['Species'],
                    y=plot_df[f"{file_name}_{energy_type}"],
                    mode='markers',
                    name=f"{file_name}",
                    marker=dict(
                        size=12,
                        symbol=symbols[j % len(symbols)],
                        color=colors,
                        line=dict(width=1, color='DarkSlateGrey')
                    ),
                    legendgroup=file_name,
                    showlegend=(energy_type == 'E'),  # Only show in legend once
                    hovertemplate=f"{file_name} %{{x}}: %{{y:.3f}} eV<extra></extra>"
                ),
                row=row, col=1
            )
        
        # Plot difference as bar plot if there are exactly 2 files
        if len(file_names) == 2:
            diff_data = plot_df[f"Diff_{energy_type}"]
            
            # Color bars by sign of difference
            bar_colors = ['rgba(255, 99, 71, 0.5)' if d >= 0 else 'rgba(60, 179, 113, 0.5)' 
                        for d in diff_data]
            
            fig.add_trace(
                go.Bar(
                    x=plot_df['Species'],
                    y=diff_data,
                    name=f"Difference",
                    marker_color=bar_colors,
                    opacity=0.6,
                    legendgroup='diff',
                    showlegend=(energy_type == 'E'),  # Only show in legend once
                    hovertemplate="Diff %{x}: %{y:.3f} eV<extra></extra>"
                ),
                row=row, col=1
            )
    
    # Add status to hover information
    for i, row in plot_df.iterrows():
        for j, energy_type in enumerate(energy_types):
            for file_name in file_names:
                value = row[f"{file_name}_{energy_type}"]
                if not pd.isna(value):
                    fig.add_annotation(
                        x=row['Species'],
                        y=value,
                        text=f"Status: {row['Status']}",
                        showarrow=False,
                        xanchor='center',
                        yanchor='bottom',
                        visible=False
                    )
    
    # Update layout
    fig.update_layout(
        height=900,
        width=1000,
        title_text="Energy Comparison Across Files",
        legend=dict(
            groupclick="toggleitem"
        ),
        hovermode="closest",
        barmode='overlay'
    )
    
    # Update yaxis titles and add grid
    for i in range(1, 4):
        fig.update_yaxes(title_text="Energy (eV)", row=i, col=1, gridcolor='lightgrey')
        fig.update_xaxes(gridcolor='lightgrey', row=i, col=1)
    
    # For bottom subplot, angle the x labels for better readability
    fig.update_xaxes(tickangle=45, row=3, col=1)
    
    # Add buttons to show/hide different energy types
    buttons = []
    for energy_type in energy_types:
        visibility = [False] * len(fig.data)
        
        # Find indices for this energy type
        indices = []
        for i, trace in enumerate(fig.data):
            if trace.legendgroup in file_names and energy_type in trace.name:
                visibility[i] = True
                indices.append(i)
            elif trace.legendgroup == 'diff' and energy_type in trace.name:
                visibility[i] = True
                indices.append(i)
                
        buttons.append(
            dict(
                label=energy_type_labels.get(energy_type, energy_type),
                method="update",
                args=[{"visible": visibility}]
            )
        )
    
    # Add button to show all
    all_visible = [True] * len(fig.data)
    buttons.append(
        dict(
            label="All",
            method="update",
            args=[{"visible": all_visible}]
        )
    )
    
    fig.update_layout(
        updatemenus=[
            dict(
                type="buttons",
                direction="right",
                buttons=buttons,
                pad={"r": 10, "t": 10},
                showactive=True,
                x=0.11,
                xanchor="left",
                y=1.1,
                yanchor="top"
            )
        ]
    )
    
    # Save if output path provided
    if output_path:
        fig.write_html(output_path)
    
    return fig

def save_gibbs_table_to_word(
    df: pd.DataFrame, 
    output_path: str,
    simple_format: bool = True
) -> None:
    """
    Save a Gibbs table to a Word document optimized for easy copy-pasting.
    
    Args:
        df: DataFrame containing the Gibbs table (can be from compare_multiple_gibbs_tables)
        output_path: Path to save the Word document
        simple_format: If True, use a simplified format
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx is not installed. Please install it with pip install python-docx")
        return
    
    document = Document()
    
    # Add title
    title = document.add_heading('Gibbs Energy Comparison Table', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a copy of the DataFrame to avoid modifying the original
    df_copy = df.copy()
    
    # Pre-process any Series objects
    for col in df_copy.columns:
        df_copy[col] = df_copy[col].apply(
            lambda x: x.iloc[0] if isinstance(x, pd.Series) and not x.empty else x
        )
    
    # Create a simpler version of the table for Word
    if simple_format:
        # Determine if we're dealing with a comparison table
        is_comparison = isinstance(df_copy.columns, pd.MultiIndex)
        
        if is_comparison:
            # For comparison tables, simplify the column structure
            # Create a flat DataFrame suitable for Word
            simple_columns = []
            file_names = set()
            
            # Get all unique file names (excluding 'Info' and 'Diff')
            for col in df_copy.columns:
                if col[0] not in ['Info', 'Diff']:
                    file_names.add(col[0])
            
            # Create new column headers
            for col in df_copy.columns:
                if col[0] == 'Info':
                    simple_columns.append(f"{col[1]}")
                elif col[0] == 'Diff':
                    simple_columns.append(f"Diff_{col[1]}")
                else:
                    simple_columns.append(f"{col[0]}_{col[1]}")
            
            # Create the simplified DataFrame
            simple_df = pd.DataFrame(columns=simple_columns)
            all_row_data = []
            
            # Fill the simplified DataFrame with values
            for i, idx in enumerate(df_copy.index):
                row_data = {}
                for j, col in enumerate(df_copy.columns):
                    value = df_copy.iloc[i, j]
                    
                    if col[0] == 'Info':
                        col_name = f"{col[1]}"
                    elif col[0] == 'Diff':
                        col_name = f"Diff_{col[1]}"
                    else:
                        col_name = f"{col[0]}_{col[1]}"
                    
                    row_data[col_name] = value
                
                all_row_data.append(row_data)
            
            # Use pd.concat() instead of append() method
            if all_row_data:
                simple_df = pd.concat([simple_df, pd.DataFrame(all_row_data)], ignore_index=True)
        else:
            # For single Gibbs tables, use specific columns
            cols_to_keep = ['species_name', 'status', 'E', 'ZPE', 'H', 'TS', 'F', 'G']
            available_cols = [col for col in cols_to_keep if col in df_copy.columns]
            simple_df = df_copy[available_cols].copy()
    else:
        simple_df = df_copy.copy()
    
    # Round numeric columns to 2 decimal places
    numeric_cols = simple_df.select_dtypes(include=['float64']).columns
    simple_df[numeric_cols] = simple_df[numeric_cols].round(2)
    
    # Create table
    table = document.add_table(rows=simple_df.shape[0] + 1, cols=simple_df.shape[1])
    table.style = 'Table Grid'
    
    # Add headers
    for j, col in enumerate(simple_df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        
        # Apply formatting to header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Add data
    for i, idx in enumerate(simple_df.index):
        for j, col in enumerate(simple_df.columns):
            value = simple_df.iloc[i, j]
            cell = table.cell(i + 1, j)
            
            # Format cell text based on value type
            if isinstance(value, pd.Series):
                # Handle Series values
                if value.empty:
                    cell.text = "N/A"
                else:
                    # Use the first value or handle as needed
                    val = value.iloc[0]
                    if pd.isna(val):
                        cell.text = "N/A"
                    elif isinstance(val, (int, float)):
                        cell.text = f"{val:.2f}"  # 2 decimal places
                    else:
                        cell.text = str(val)
            elif pd.isna(value):
                cell.text = "N/A"
            elif isinstance(value, (int, float)):
                cell.text = f"{value:.2f}"  # 2 decimal places
            else:
                cell.text = str(value)
    
    # Add a description
    document.add_paragraph('')
    
    # Add a legend for the values
    legend_text = 'E: Formation Energy, ZPE: Zero-Point Energy, H: Enthalpy, TS: Temperature×Entropy, F: Helmholtz Free Energy, G: Gibbs Free Energy'
    
    # If it's a comparison table, add more explanation
    if isinstance(df.columns, pd.MultiIndex):
        legend_text += '\nDiff_X: Difference in value X between the compared files'
        
    document.add_paragraph(legend_text)
    
    # Save the document
    document.save(output_path)
    print(f"Word document saved to {output_path}")

# Add a simple wrapper function to save the comparison table
def save_comparison_to_word(
    comparison_df: pd.DataFrame,
    output_path: str
) -> None:
    """
    Save a comparison table to a Word document for easy copy-pasting.
    
    Args:
        comparison_df: DataFrame from compare_multiple_gibbs_tables
        output_path: Path to save the Word document
    """
    # Make sure the directory exists
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Add .docx extension if not present
    if output_path.suffix.lower() != '.docx':
        output_path = output_path.with_suffix('.docx')
    
    save_gibbs_table_to_word(comparison_df, output_path)

def save_gibbs_table_to_latex(
    df: pd.DataFrame, 
    output_path: str,
    include_diff: bool = True
) -> None:
    """
    Save Gibbs energy comparison table to LaTeX format (.tex)
    
    Parameters
    ----------
    df : pd.DataFrame
        DataFrame from compare_multiple_gibbs_tables
    output_path : str
        Path to save the LaTeX file
    include_diff : bool, optional
        Whether to include difference columns, by default True
    """
    # Convert output_path to Path object
    output_path = Path(output_path)
    
    # Ensure the path has .tex extension
    if output_path.suffix != '.tex':
        output_path = output_path.with_suffix('.tex')
        
    # Create directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Extract file names (excluding 'Info' and optionally 'Diff')
    file_names = sorted(set([col[0] for col in df.columns 
                       if col[0] not in ['Info'] + ([] if include_diff else ['Diff'])]))
    
    # Extract column types (E, F, G)
    col_types = sorted(set([col[1] for col in df.columns 
                      if col[0] != 'Info' and col[1] not in ['surface_name', 'site_name', 'species_name', 'status']]))
    
    # Prepare LaTeX code
    latex_code = []
    latex_code.append("\\documentclass{article}")
    latex_code.append("\\usepackage{booktabs}")
    latex_code.append("\\usepackage{multirow}")
    latex_code.append("\\usepackage{array}")
    latex_code.append("\\usepackage{siunitx}")
    latex_code.append("\\usepackage[table]{xcolor}")
    latex_code.append("\\begin{document}")
    
    # Table header
    latex_code.append("\\begin{table}")
    latex_code.append("\\centering")
    latex_code.append("\\caption{Gibbs Energy Comparison}")
    
    # Calculate column specifications
    column_spec = "l" * 2  # species_name and status
    file_count = len(file_names)
    for _ in range(file_count * len(col_types)):
        column_spec += "S"  # Numeric columns using siunitx
    if include_diff:
        column_spec += "S" * len(col_types)  # Diff columns
    
    latex_code.append(f"\\begin{{tabular}}{{@{{}}{column_spec}@{{}}}}")
    latex_code.append("\\toprule")
    
    # Create multicolumn header
    header_row1 = ["\\multirow{2}{*}{Species}", "\\multirow{2}{*}{Status}"]
    for file_name in file_names:
        header_row1.append(f"\\multicolumn{{{len(col_types)}}}{{c}}{{{file_name}}}")
    if include_diff:
        header_row1.append(f"\\multicolumn{{{len(col_types)}}}{{c}}{{Diff}}")
    latex_code.append(" & ".join(header_row1) + " \\\\")
    
    # Second header row with column types
    header_row2 = ["", ""]
    for _ in range(file_count + (1 if include_diff else 0)):
        header_row2.extend(col_types)
    latex_code.append(" & ".join(header_row2) + " \\\\")
    latex_code.append("\\midrule")
    
    # Table data
    for _, row in df.iterrows():
        data_row = []
        # Add species and status
        species = row[('Info', 'species_name')]
        status = row[('Info', 'status')]
        data_row.append(str(species))
        data_row.append(str(status) if not pd.isna(status) else "N/A")
        
        # Add data for each file
        for file_name in file_names:
            for col_type in col_types:
                value = row.get((file_name, col_type))
                if pd.isna(value):
                    data_row.append("N/A")
                else:
                    data_row.append(f"{value:.3f}")
        
        # Add diff columns if requested
        if include_diff:
            for col_type in col_types:
                value = row.get(('Diff', col_type))
                if pd.isna(value):
                    data_row.append("N/A")
                else:
                    # Highlight differences
                    threshold = df.attrs.get('threshold', 0.01)
                    if abs(value) > threshold:
                        color = 'red' if value > 0 else 'green'
                        data_row.append(f"\\textcolor{{{color}}}{{{value:.3f}}}")
                    else:
                        data_row.append(f"{value:.3f}")
        
        latex_code.append(" & ".join(data_row) + " \\\\")
    
    # Table footer
    latex_code.append("\\bottomrule")
    latex_code.append("\\end{tabular}")
    latex_code.append("\\end{table}")
    latex_code.append("\\end{document}")
    
    # Write to file
    with open(output_path, 'w') as f:
        f.write("\n".join(latex_code))

def create_scatter_plots(
    comparison_df, 
    output_path: str = None,
    dpi: int = 300,
    figsize: tuple = (12, 15)
):
    """
    Create and save scatter plots of energy values with differences
    
    Parameters
    ----------
    comparison_df : pd.DataFrame
        DataFrame from compare_multiple_gibbs_tables
    output_path : str, optional
        Path to save the PNG file, by default None
    dpi : int, optional
        Resolution for saved image, by default 300
    figsize : tuple, optional
        Figure size, by default (12, 15)
        
    Returns
    -------
    matplotlib.figure.Figure
        The created figure
    """
    import matplotlib.pyplot as plt
    import numpy as np
    
    # Extract file names (excluding 'Info' and 'Diff')
    file_names = [col[0] for col in comparison_df.columns 
                 if col[0] not in ['Info', 'Diff']]
    
    # Energy types to plot
    energy_types = ['E', 'F', 'G']
    
    # Create figure
    fig, axs = plt.subplots(3, 1, figsize=figsize, sharex=True)
    
    # Set species names as x-axis values
    species_names = comparison_df[('Info', 'species_name')]
    status = comparison_df[('Info', 'status')]
    
    # Color-code by status
    colors = {'gas': 'tab:red', 'ads': 'tab:blue'}
    default_color = 'tab:gray'
    
    # For each energy type
    for i, energy_type in enumerate(energy_types):
        ax = axs[i]
        
        # Plot data for each file
        markers = ['o', 's', '^', 'D', 'v', '*']  # Different markers for different files
        
        for j, file_name in enumerate(file_names):
            # Extract data
            y_data = comparison_df[(file_name, energy_type)]
            
            # Determine colors based on status
            point_colors = [colors.get(str(s).lower(), default_color) for s in status]
            
            # Plot data points
            ax.scatter(range(len(species_names)), y_data, 
                      marker=markers[j % len(markers)], 
                      color=point_colors,
                      s=80, alpha=0.7, label=f"{file_name}")
        
        # Plot difference as bar plot if there are exactly 2 files
        if len(file_names) == 2 and 'Diff' in [col[0] for col in comparison_df.columns]:
            diff_data = comparison_df[('Diff', energy_type)]
            bar_colors = ['indianred' if d >= 0 else 'mediumseagreen' for d in diff_data]
            
            # Plot difference bars
            diff_bars = ax.bar(range(len(species_names)), diff_data, 
                              alpha=0.3, color=bar_colors, width=0.5,
                              label='Difference')
            
            # Add secondary y-axis for difference
            ax2 = ax.twinx()
            ax2.set_ylabel('Difference (eV)', color='gray')
            ax2.tick_params(axis='y', labelcolor='gray')
            ax2.set_ylim(diff_data.min() * 1.5 if not pd.isna(diff_data.min()) else -0.2, 
                        diff_data.max() * 1.5 if not pd.isna(diff_data.max()) else 0.2)
            
            # Set visible to False to avoid duplicate in legend
            for bar in diff_bars:
                bar.set_visible(False)
        
        # Set titles and labels
        ax.set_title(f"{energy_type_labels.get(energy_type, energy_type)}")
        ax.set_ylabel('Energy (eV)')
        ax.grid(True, alpha=0.3)
        
        # Only add legend to first plot to avoid duplication
        if i == 0:
            handles, labels = ax.get_legend_handles_labels()
            if len(file_names) == 2:
                # Add manual entry for difference bars
                from matplotlib.patches import Patch
                handles.append(Patch(facecolor='indianred', alpha=0.3, label='Diff (>0)'))
                handles.append(Patch(facecolor='mediumseagreen', alpha=0.3, label='Diff (<0)'))
                
            ax.legend(handles=handles, loc='best', frameon=True, 
                     framealpha=0.9, ncol=min(3, len(handles)))
    
    # Set x-ticks to species names
    axs[-1].set_xticks(range(len(species_names)))
    axs[-1].set_xticklabels(species_names, rotation=90)
    
    # Add color legend for status
    from matplotlib.patches import Patch
    status_legend_elements = [Patch(facecolor=color, label=s.upper()) 
                             for s, color in colors.items() 
                             if s in [str(st).lower() for st in status.unique()]]
    fig.legend(handles=status_legend_elements, loc='lower center', 
              bbox_to_anchor=(0.5, 0), ncol=len(status_legend_elements))
    
    # Adjust layout
    plt.tight_layout()
    plt.subplots_adjust(hspace=0.3, bottom=0.15)
    
    # Save if output path provided
    if output_path:
        plt.savefig(output_path, dpi=dpi, bbox_inches='tight')
    
    return fig

# Dictionary mapping energy types to labels
energy_type_labels = {
    'E': 'Formation Energy (E)',
    'F': 'Free Energy Correction (F)',
    'G': 'Gibbs Energy (G)'
}